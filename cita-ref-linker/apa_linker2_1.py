import re
import time
import sys
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- ANIMATION UTILS ---
def print_progress_bar(iteration, total, prefix='', suffix='', decimals=1, length=40, fill='█', printEnd="\r"):
    percent = ("{0:." + str(decimals) + "f}").format(100 * (iteration / float(total)))
    filled_length = int(length * iteration // total)
    bar = fill * filled_length + '-' * (length - filled_length)
    sys.stdout.write(f'\r{prefix} |{bar}| {percent}% {suffix}')
    sys.stdout.flush()
    if iteration == total: 
        print()

def fake_thinking_time():
    time.sleep(0.005) 

# --- XML HELPER FUNCTIONS ---
def add_bookmark(p, bookmark_name, bookmark_id):
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bookmark_id))
    start.set(qn('w:name'), bookmark_name)
    p._p.insert(0, start)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bookmark_id))
    p._p.append(end)

def create_hyperlink_run(paragraph, text, bookmark_name, font_name=None, font_size=None):
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)
    hyperlink.set(qn('w:history'), '1')

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(color)
    rPr.append(underline)

    if font_name:
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rPr.append(rFonts)
    if font_size:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(font_size.pt * 2)))
        rPr.append(sz)

    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    hyperlink.append(r)
    paragraph._p.append(hyperlink)

def clean_author_name(raw_text):
    """
    Extract surname only.
    UPDATED: Now supports Unicode characters (ó, é, ñ) and hyphens.
    """
    first_word = raw_text.split()[0]
    # \w matches any unicode letter. We also allow hyphens and apostrophes.
    return re.sub(r"[^\w\-\']", "", first_word)

# --------------------------------------------------------
# MAIN EXECUTION
# --------------------------------------------------------

print("\n" + "="*50)
print("      AUTO-LINKER & VALIDATOR v2.1 (UNICODE)      ")
print("="*50 + "\n")

input_filename = "kusniawati_et_al.2025.docx"
print(f"[*] Loading {input_filename}...")
doc = Document(input_filename)

all_paragraphs = list(doc.paragraphs)
total_paras = len(all_paragraphs)

ref_map = {}
linked_references = set()
missing_citations = []
unique_id_counter = 0
in_refs_section = False

# UPDATED REGEX: 
# ^([\w\-\']+) -> Matches "Cantón", "O'Neil", "Smith-Jones" (stops at space/comma)
ref_pattern = r"^([\w\-\']+).*?\(?(\d{4})\)?"

# --- PHASE 1: MAPPING ---
print("\n[*] Tahap 1: Menscan Daftar Referensi")

for i, p in enumerate(all_paragraphs):
    print_progress_bar(i + 1, total_paras, prefix='Scanning:', suffix='Selesai', length=30)
    fake_thinking_time()

    if "References" in p.text and len(p.text) < 50:
        in_refs_section = True
        continue
    
    if in_refs_section:
        match = re.search(ref_pattern, p.text.strip())
        if match:
            surname = match.group(1)
            year = match.group(2)
            key = f"{clean_author_name(surname)}_{year}"
            
            # Sanitize bookmark name (XML doesn't like accents in ID names sometimes)
            safe_key = re.sub(r"[^A-Za-z0-9]", "", key) 
            bookmark_name = f"REF_{safe_key}_{unique_id_counter}"
            
            add_bookmark(p, bookmark_name, unique_id_counter)
            ref_map[key] = bookmark_name
            unique_id_counter += 1

print(f"    > Found {len(ref_map)} unique references.")

# --- PHASE 2: LINKING ---
print("\n[*] Tahap 2: Mengelink Sitasi")
citation_pattern = r"\((.*?)\)"

for i, p in enumerate(all_paragraphs):
    print_progress_bar(i + 1, total_paras, prefix='Mengelink :', suffix='Selesai', length=30)
    fake_thinking_time()

    if "References" in p.text and len(p.text) < 50:
        break

    text = p.text
    if "(" not in text:
        continue

    matches = list(re.finditer(citation_pattern, text))
    if not matches:
        continue

    original_font_name = p.runs[0].font.name if p.runs else None
    original_font_size = p.runs[0].font.size if p.runs else None

    p.text = "" 
    cursor = 0
    
    for match in matches:
        content = match.group(1)
        start_index = match.start()
        if start_index > cursor:
            run = p.add_run(text[cursor:start_index])
            if original_font_name: run.font.name = original_font_name
            if original_font_size: run.font.size = original_font_size
            
        run = p.add_run("(")
        if original_font_name: run.font.name = original_font_name
        if original_font_size: run.font.size = original_font_size
        
        citations_in_group = content.split(";")
        
        for k, cite_text in enumerate(citations_in_group):
            cite_text = cite_text.strip()
            
            # UPDATED: Use \w in citation matching too
            sub_match = re.search(r"(.*),\s.*?(\d{4})", cite_text)
            
            if sub_match:
                raw_author = sub_match.group(1)
                year = sub_match.group(2)
                search_key = f"{clean_author_name(raw_author)}_{year}"
                
                if search_key in ref_map:
                    linked_references.add(search_key)
                    create_hyperlink_run(
                        p, cite_text, ref_map[search_key], 
                        font_name=original_font_name, font_size=original_font_size
                    )
                else:
                    missing_citations.append(cite_text)
                    run = p.add_run(cite_text)
                    if original_font_name: run.font.name = original_font_name
                    if original_font_size: run.font.size = original_font_size
            else:
                run = p.add_run(cite_text)
                if original_font_name: run.font.name = original_font_name
                if original_font_size: run.font.size = original_font_size
                
            if k < len(citations_in_group) - 1:
                run = p.add_run("; ")
                if original_font_name: run.font.name = original_font_name
                if original_font_size: run.font.size = original_font_size
        
        run = p.add_run(")")
        if original_font_name: run.font.name = original_font_name
        if original_font_size: run.font.size = original_font_size
        cursor = match.end()

    if cursor < len(text):
        run = p.add_run(text[cursor:])
        if original_font_name: run.font.name = original_font_name
        if original_font_size: run.font.size = original_font_size

# --- REPORTING ---
print("\n[*] Menyimpan Dokumen...")
doc.save("output_kusniawati_linked_v2.docx")

all_ref_keys = set(ref_map.keys())
unused_references = all_ref_keys - linked_references

print("[*] Membuat Laporan...")
time.sleep(1)

print("\n" + "="*40)
print(f" PROSES SELESAI ")
print("="*40)
print(f" [+] Total Referensi : {len(all_ref_keys)}")
print(f" [+] Sitasi Terhubung : {len(linked_references)}")
print(f" [-] Sitasi Rusak : {len(missing_citations)}")
print(f" [-] Referensi Bermasalah: {len(unused_references)}")
print("="*40 + "\n")

with open("validation_report.txt", "w", encoding="utf-8") as f:
    f.write("LAPORAN VALIDASI SITASI\n")
    f.write("==========================\n\n")
    f.write(f"SITASI RUSAK ({len(missing_citations)}):\n")
    if missing_citations:
        for c in sorted(list(set(missing_citations))): f.write(f" [x] {c}\n")
    else: f.write(" - None\n")
    
    f.write(f"\nREFERENSI RUSAK ({len(unused_references)}):\n")
    if unused_references:
        for r in sorted(list(unused_references)): f.write(f" [?] {r.replace('_', ', ')}\n")
    else: f.write(" - None\n")