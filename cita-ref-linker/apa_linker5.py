import re
import time
import sys
import os  # <--- NEW: Needed for folder creation
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- ANIMATION UTILS ---
def print_progress_bar(iteration, total, prefix='', suffix='', decimals=1, length=40, fill='█'):
    percent = ("{0:." + str(decimals) + "f}").format(100 * (iteration / float(total)))
    filled_length = int(length * iteration // total)
    bar = fill * filled_length + '-' * (length - filled_length)
    sys.stdout.write(f'\r{prefix} |{bar}| {percent}% {suffix}')
    sys.stdout.flush()
    if iteration == total: 
        print()

def fake_thinking_time():
    time.sleep(0.002) 

# --- XML HELPER FUNCTIONS ---
def add_bookmark(p, bookmark_name, bookmark_id):
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bookmark_id))
    start.set(qn('w:name'), bookmark_name)
    p._p.insert(0, start)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bookmark_id))
    p._p.append(end)

def create_hyperlink_run(paragraph, text, bookmark_name, font_name="Calibri", font_size=None):
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)
    hyperlink.set(qn('w:history'), '1')

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(color)
    rPr.append(underline)

    # Force Font Preservation (Default to Calibri)
    actual_font = font_name if font_name else "Calibri"
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), actual_font)
    rFonts.set(qn('w:hAnsi'), actual_font)
    rPr.append(rFonts)

    if font_size:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(font_size.pt * 2)))
        rPr.append(sz)

    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    hyperlink.append(r)
    paragraph._p.append(hyperlink)

def clean_author_name(raw_text):
    first_word = raw_text.split()[0] 
    return re.sub(r"[^\w\-\']", "", first_word)

# --------------------------------------------------------
# MAIN EXECUTION
# --------------------------------------------------------

print("\n" + "="*50)
print("      AUTO-LINKER v5.0 (INTERACTIVE)      ")
print("="*50 + "\n")

# --- NEW: INTERACTIVE INPUT ---
while True:
    input_filename = input("Enter the filename (e.g. data/file.docx): ").strip()
    # Remove quotes if user dragged and dropped file
    input_filename = input_filename.replace('"', '').replace("'", "")
    
    if os.path.exists(input_filename):
        break
    else:
        print(f"[!] File not found: {input_filename}")
        print("    Please try again.\n")

# --- NEW: FOLDER SETUP ---
# Extract "kusniawati_et_al.2025" from "kusniawati_et_al.2025.docx"
base_name = os.path.splitext(os.path.basename(input_filename))[0]
output_folder = base_name  # The folder name will be the file name

# Create the folder if it doesn't exist
if not os.path.exists(output_folder):
    os.makedirs(output_folder)
    print(f"[*] Created output folder: {output_folder}/")
else:
    print(f"[*] Using existing folder: {output_folder}/")

print(f"[*] Loading {input_filename}...")
doc = Document(input_filename)

all_paragraphs = list(doc.paragraphs)
total_paras = len(all_paragraphs)

ref_map = {}
linked_references = set()
missing_citations = []
unique_id_counter = 0
in_refs_section = False

ref_list_pattern = r"^([\w\-\']+).*?\(?(\d{4})\)?"

# --- PHASE 1: MAPPING ---
print("\n[*] Phase 1: Mapping References")

for i, p in enumerate(all_paragraphs):
    print_progress_bar(i + 1, total_paras, prefix='Scanning:', suffix='Done', length=30)
    fake_thinking_time()

    if "References" in p.text and len(p.text) < 50:
        in_refs_section = True
        continue
    
    if in_refs_section:
        match = re.search(ref_list_pattern, p.text.strip())
        if match:
            surname = match.group(1)
            year = match.group(2)
            key = f"{clean_author_name(surname)}_{year}"
            
            safe_key = re.sub(r"[^A-Za-z0-9]", "", key) 
            bookmark_name = f"REF_{safe_key}_{unique_id_counter}"
            
            add_bookmark(p, bookmark_name, unique_id_counter)
            ref_map[key] = bookmark_name
            unique_id_counter += 1

print(f"    > Mapped {len(ref_map)} references.")

# --- PHASE 2: LINKING ---
print("\n[*] Phase 2: Linking Citations")

citation_pattern = re.compile(
    r"(?P<paren>\([^\)]+\))|" 
    r"(?P<narrative>[A-Z][\w\-\']+(?:\s+(?:&|and)\s+[A-Z][\w\-\']+)?(?:\s+et\s+al\.?)?\s*\(\d{4}\))"
)

for i, p in enumerate(all_paragraphs):
    print_progress_bar(i + 1, total_paras, prefix='Linking :', suffix='Done', length=30)
    fake_thinking_time()

    if "References" in p.text and len(p.text) < 50:
        break

    text = p.text
    if "(" not in text: 
        continue

    matches = list(citation_pattern.finditer(text))
    if not matches:
        continue

    original_font_name = None
    original_font_size = None
    
    if p.runs:
        for run in p.runs:
            if run.font.name:
                original_font_name = run.font.name
                break 
        original_font_size = p.runs[0].font.size

    if not original_font_name:
        original_font_name = "Calibri"

    p.text = "" 
    cursor = 0
    
    for match in matches:
        full_text = match.group(0)
        start_index = match.start()
        
        if start_index > cursor:
            run = p.add_run(text[cursor:start_index])
            run.font.name = original_font_name
            if original_font_size: run.font.size = original_font_size
        
        # LINKING LOGIC
        if match.group('narrative'):
            parts = full_text.split('(')
            name_part = parts[0].strip()
            year_part = parts[1].replace(')', '').strip()
            key = f"{clean_author_name(name_part)}_{year_part}"
            
            if key in ref_map:
                linked_references.add(key)
                create_hyperlink_run(p, full_text, ref_map[key], original_font_name, original_font_size)
            else:
                missing_citations.append(full_text)
                run = p.add_run(full_text)
                run.font.name = original_font_name
                if original_font_size: run.font.size = original_font_size

        elif match.group('paren'):
            content = full_text[1:-1]
            run = p.add_run("(")
            run.font.name = original_font_name
            if original_font_size: run.font.size = original_font_size
            
            sub_cites = content.split(";")
            for k, cite in enumerate(sub_cites):
                cite = cite.strip()
                sub_match = re.search(r"(.*),\s.*?(\d{4})", cite)
                
                if sub_match:
                    raw_author = sub_match.group(1)
                    year = sub_match.group(2)
                    key = f"{clean_author_name(raw_author)}_{year}"
                    
                    if key in ref_map:
                        linked_references.add(key)
                        create_hyperlink_run(p, cite, ref_map[key], original_font_name, original_font_size)
                    else:
                        missing_citations.append(cite)
                        run = p.add_run(cite)
                        run.font.name = original_font_name
                        if original_font_size: run.font.size = original_font_size
                else:
                    run = p.add_run(cite)
                    run.font.name = original_font_name
                    if original_font_size: run.font.size = original_font_size
                
                if k < len(sub_cites) - 1:
                    run = p.add_run("; ")
                    run.font.name = original_font_name
                    if original_font_size: run.font.size = original_font_size
            
            run = p.add_run(")")
            run.font.name = original_font_name
            if original_font_size: run.font.size = original_font_size

        cursor = match.end()

    if cursor < len(text):
        run = p.add_run(text[cursor:])
        run.font.name = original_font_name
        if original_font_size: run.font.size = original_font_size

# --- NEW: SAVE TO FOLDER ---
output_doc_path = os.path.join(output_folder, f"{base_name}_linked.docx")
print(f"\n[*] Saving Document to: {output_doc_path}")
doc.save(output_doc_path)

# --- NEW: REPORT TO FOLDER ---
output_report_path = os.path.join(output_folder, "validation_report.txt")
print(f"[*] Generating Report to: {output_report_path}")

all_ref_keys = set(ref_map.keys())
unused_references = all_ref_keys - linked_references

with open(output_report_path, "w", encoding="utf-8") as f:
    f.write(f"VALIDATION REPORT FOR: {input_filename}\n")
    f.write("="*50 + "\n\n")
    f.write(f"BROKEN CITATIONS ({len(missing_citations)}):\n")
    for c in sorted(list(set(missing_citations))): f.write(f" [x] {c}\n")
    
    f.write(f"\nUNUSED REFERENCES ({len(unused_references)}):\n")
    for r in sorted(list(unused_references)): f.write(f" [?] {r}\n")

print("\n" + "="*40)
print(f" JOB DONE! CHECK FOLDER: {output_folder}/ ")
print("="*40 + "\n")