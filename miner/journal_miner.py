import re
from docx import Document
from docxtpl import DocxTemplate

def extract_metadata_and_body(raw_filename):
    """
    1. Extracts simple strings for Title/Abstract (metadata).
    2. Extracts the 'Body' as a sub-document to preserve italics/tables.
    """
    doc = Document(raw_filename)
    data = {}
    
    # --- PART A: HEURISTIC MINING (Get the Strings) ---
    # We assume standard order: Title -> Author -> Affiliation -> Abstract -> Body
    
    # 1. Get Title (First non-empty paragraph)
    paras = [p for p in doc.paragraphs if p.text.strip()]
    data['title'] = paras[0].text
    data['author'] = paras[1].text
    data['affiliation'] = paras[2].text
    
    # 2. Get Abstract & Keywords
    # We scan for the word "Abstract"
    data['abstract'] = "Abstract not found."
    data['keywords'] = ""
    
    body_start_index = 0
    
    for i, p in enumerate(paras):
        text = p.text.strip()
        if text.lower() == 'abstract':
            # The *next* paragraph is usually the abstract content
            data['abstract'] = paras[i+1].text
        
        if text.lower().startswith('keywords:'):
            data['keywords'] = text.replace('Keywords:', '').strip()
            
        # Stop 'metadata mining' once we hit Introduction
        if text.lower() in ['introduction', 'background']:
            body_start_index = i # Mark where the real article starts
            break
            
    # --- PART B: THE BODY TRANSPLANT (Preserve Formatting) ---
    # We create a new "blank" document that will hold ONLY the body.
    # This preserves italics, bold, links, etc. from the original.
    
    body_doc = Document()
    
    # Copy paragraphs starting from "Introduction" to the end
    # Note: This is a simplified copy. For full fidelity (images/tables), 
    # we would use a more advanced 'element move', but this works for text/italics.
    
    source_paras = doc.paragraphs  # All paragraphs in source
    
    # Find the index in the *full* paragraph list where body starts
    real_start_idx = 0
    for idx, p in enumerate(source_paras):
        if p.text.strip().lower() in ['introduction', 'background']:
            real_start_idx = idx
            break
            
    # Copy the content into our 'body_doc'
    for p in source_paras[real_start_idx:]:
        # Create a new paragraph in the destination
        new_p = body_doc.add_paragraph()
        new_p.style = p.style # Copy style (Normal, Heading 1, etc.)
        
        # Copy runs (this preserves BOLD and ITALIC)
        for run in p.runs:
            new_run = new_p.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            
    return data, body_doc

# ==========================================
# MAIN EXECUTION
# ==========================================

# 1. Setup Files
raw_file = "data/palupi.docx"      # The input submission
template_file = "data/journal_template.docx" # The template you made manually

print(f"[*] Mining {raw_file}...")
meta_data, body_subdoc = extract_metadata_and_body(raw_file)

print(f"    Title: {meta_data['title'][:30]}...")
print(f"    Author: {meta_data['author']}")

# 2. Load Template
print(f"[*] Loading Template: {template_file}...")
tpl = DocxTemplate(template_file)

# 3. Handle the "Sub-Document"
# We tell the template: "This variable is not text, it's a DOCX file"
sd = tpl.new_subdoc()
sd.element.body.extend(body_subdoc.element.body) # Low-level XML copy

# 4. Prepare Context
context = {
    'title': meta_data['title'],
    'author': meta_data['author'],
    'abstract': meta_data['abstract'],
    'keywords': meta_data['keywords'],
    'body_content': sd,  # Inject the sub-document here
}

# 5. Render and Save
print("[*] Rendering Final Journal...")
tpl.render(context)
tpl.save("output/PUBLICATION_READY.docx")
print("DONE. Check output/PUBLICATION_READY.docx")