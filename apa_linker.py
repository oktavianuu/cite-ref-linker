import re
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- XML HELPER FUNCTIONS ---
def add_bookmark(p, bookmark_name, bookmark_id):
    """Add an invisible bookmark to the paragraph."""
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bookmark_id))
    start.set(qn('w:name'), bookmark_name)
    p._p.insert(0, start)

    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bookmark_id))
    p._p.append(end)

def create_hyperlink_run(paragraph, text, bookmark_name, font_name=None, font_size=None):
    """Appends a clickable hyperlink with preserved font settings."""
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)
    hyperlink.set(qn('w:history'), '1')

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Blue + Underline
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(color)
    rPr.append(underline)

    # Restore Font
    if font_name:
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rPr.append(rFonts)
    if font_size:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(font_size.pt * 2)))
        rPr.append(sz)

    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    hyperlink.append(r)
    paragraph._p.append(hyperlink)

def clean_author_name(raw_text):
    """Extract surname only."""
    first_word = raw_text.split()[0]
    return re.sub(r"[^A-Za-z]", "", first_word)

# --------------------------------------------------------
# MAIN EXECUTION
# --------------------------------------------------------

input_filename = "palupi.docx"
doc = Document(input_filename)

# Tracking Data Structures
ref_map = {}            # Stores { "Thompson_2025": "bookmark_name" }
linked_references = set() # Stores keys that were successfully linked
missing_citations = []    # Stores text of citations that failed to link

unique_id_counter = 0
in_refs_section = False
ref_pattern = r"^([A-Za-z\-]+).*?\((\d{4})\)"

print("--- PHASE 1: Scanning Reference List ---")

for p in doc.paragraphs:
    if "References" in p.text and len(p.text) < 50:
        in_refs_section = True
        continue
    
    if in_refs_section:
        match = re.search(ref_pattern, p.text.strip())
        if match:
            surname = match.group(1)
            year = match.group(2)
            key = f"{clean_author_name(surname)}_{year}"
            
            bookmark_name = f"REF_{key}_{unique_id_counter}"
            add_bookmark(p, bookmark_name, unique_id_counter)
            
            ref_map[key] = bookmark_name
            unique_id_counter += 1

print(f"Mapped {len(ref_map)} references in the bibliography.")

# --------------------------------------------------------
# PHASE 2: LINKING & CHECKING
# --------------------------------------------------------
print("--- PHASE 2: Linking and Validating Citations ---")
citation_pattern = r"\((.*?)\)"

for p in doc.paragraphs:
    # Stop processing if we hit the references section
    if "References" in p.text and len(p.text) < 50:
        break

    text = p.text
    if "(" not in text:
        continue

    matches = list(re.finditer(citation_pattern, text))
    if not matches:
        continue

    # Capture original font formatting
    original_font_name = p.runs[0].font.name if p.runs else None
    original_font_size = p.runs[0].font.size if p.runs else None

    # Rebuild Paragraph
    p.text = "" 
    cursor = 0
    
    for match in matches:
        content = match.group(1)
        
        # Add text BEFORE citation
        start_index = match.start()
        if start_index > cursor:
            run = p.add_run(text[cursor:start_index])
            if original_font_name: run.font.name = original_font_name
            if original_font_size: run.font.size = original_font_size
            
        # Add Opening Parenthesis
        run = p.add_run("(")
        if original_font_name: run.font.name = original_font_name
        if original_font_size: run.font.size = original_font_size
        
        citations_in_group = content.split(";")
        
        for i, cite_text in enumerate(citations_in_group):
            cite_text = cite_text.strip()
            
            # Extract Author and Year from citation text
            sub_match = re.search(r"(.*),\s.*?(\d{4})", cite_text)
            
            if sub_match:
                raw_author = sub_match.group(1)
                year = sub_match.group(2)
                search_key = f"{clean_author_name(raw_author)}_{year}"
                
                if search_key in ref_map:
                    # VALID CITATION: Link it and track it
                    linked_references.add(search_key)
                    
                    create_hyperlink_run(
                        p, cite_text, ref_map[search_key], 
                        font_name=original_font_name, font_size=original_font_size
                    )
                else:
                    # INVALID CITATION: Track it for the report
                    missing_citations.append(cite_text)
                    
                    # Add as normal text (no blue link)
                    run = p.add_run(cite_text)
                    if original_font_name: run.font.name = original_font_name
                    if original_font_size: run.font.size = original_font_size
            else:
                # Not a standard citation format (e.g. "see Fig 1")
                run = p.add_run(cite_text)
                if original_font_name: run.font.name = original_font_name
                if original_font_size: run.font.size = original_font_size
                
            if i < len(citations_in_group) - 1:
                run = p.add_run("; ")
                if original_font_name: run.font.name = original_font_name
                if original_font_size: run.font.size = original_font_size
        
        # Add Closing Parenthesis
        run = p.add_run(")")
        if original_font_name: run.font.name = original_font_name
        if original_font_size: run.font.size = original_font_size
        
        cursor = match.end()

    # Add remaining text
    if cursor < len(text):
        run = p.add_run(text[cursor:])
        if original_font_name: run.font.name = original_font_name
        if original_font_size: run.font.size = original_font_size

# --------------------------------------------------------
# PHASE 3: REPORTING
# --------------------------------------------------------
doc.save("output_palupi_linked_formatted.docx")

# Calculate Unused References
# All References - Linked References = Unused References
all_ref_keys = set(ref_map.keys())
unused_references = all_ref_keys - linked_references

print(f"\nGenerating Report...")
print(f"Broken Citations: {len(missing_citations)}")
print(f"Unused References: {len(unused_references)}")

with open("validation_report.txt", "w", encoding="utf-8") as f:
    f.write("========================================\n")
    f.write("      CITATION VALIDATION REPORT        \n")
    f.write("========================================\n\n")
    
    f.write(f"TOTAL REFERENCES FOUND: {len(all_ref_keys)}\n")
    f.write(f"TOTAL CITATIONS LINKED: {len(linked_references)}\n\n")
    
    f.write("----------------------------------------\n")
    f.write("[!] BROKEN CITATIONS (In text, but not in Ref list)\n")
    f.write("----------------------------------------\n")
    if missing_citations:
        # Use set to remove duplicates, then sort
        for cite in sorted(list(set(missing_citations))):
            f.write(f" - {cite}\n")
    else:
        f.write(" (None - All citations are valid)\n")
        
    f.write("\n")
    f.write("----------------------------------------\n")
    f.write("[!] UNUSED REFERENCES (In Ref list, but never cited)\n")
    f.write("----------------------------------------\n")
    if unused_references:
        for ref_key in sorted(list(unused_references)):
            # Convert key "Smith_2020" back to readable "Smith (2020)"
            readable = ref_key.replace("_", ", ")
            f.write(f" - {readable}\n")
    else:
        f.write(" (None - All references are cited)\n")

print("DONE. Check 'validation_report.txt' for details.")